"""Text extraction helpers for local attachment files."""

from __future__ import annotations

from dataclasses import dataclass
import mimetypes
from pathlib import Path

from personal_agent.text_safety import clean_text


TEXT_SUFFIXES = {
    ".cfg",
    ".csv",
    ".ini",
    ".json",
    ".log",
    ".md",
    ".py",
    ".rst",
    ".toml",
    ".txt",
    ".yaml",
    ".yml",
}


@dataclass(frozen=True)
class ExtractedAttachmentText:
    text: str
    source_type: str
    truncated: bool = False
    chars: int = 0
    original_chars: int = 0
    pages_read: int = 0


class AttachmentTextExtractError(RuntimeError):
    def __init__(self, reason: str, detail: str = "") -> None:
        self.reason = reason
        self.detail = detail
        super().__init__(detail or reason)


def extract_attachment_text(
    path: str | Path,
    *,
    mime_type: str = "",
    name: str = "",
    max_chars: int = 12000,
    pdf_max_pages: int = 20,
) -> ExtractedAttachmentText:
    target = Path(path)
    if not target.exists() or not target.is_file():
        raise AttachmentTextExtractError("file_not_found")

    max_chars = max(1, int(max_chars or 12000))
    pdf_max_pages = max(1, int(pdf_max_pages or 20))
    suffix = _suffix(target, name)
    mime_type = _mime_type(target, mime_type, name)

    if suffix == ".pdf" or mime_type == "application/pdf":
        text, pages_read = _read_pdf(target, pdf_max_pages=pdf_max_pages)
        return _finish(text, "pdf", max_chars=max_chars, pages_read=pages_read)
    if suffix == ".docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _finish(_read_docx(target), "docx", max_chars=max_chars)
    if suffix in TEXT_SUFFIXES or mime_type.startswith("text/"):
        return _finish(_read_plain_text(target), "text", max_chars=max_chars)

    raise AttachmentTextExtractError("unsupported_file_type", suffix or mime_type or "unknown")


def _read_plain_text(path: Path) -> str:
    data = path.read_bytes()
    if b"\x00" in data[:4096]:
        raise AttachmentTextExtractError("unsupported_file_type", "binary-like text file")
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_pdf(path: Path, *, pdf_max_pages: int) -> tuple[str, int]:
    try:
        import fitz  # pymupdf
    except Exception as exc:  # pragma: no cover - dependency is installed in project env
        raise AttachmentTextExtractError("text_extract_unavailable", str(exc)) from exc

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        raise AttachmentTextExtractError("text_extract_failed", str(exc)) from exc
    try:
        page_count = min(len(doc), pdf_max_pages)
        text = "\n\n".join(doc[index].get_text() for index in range(page_count))
        return text, page_count
    except Exception as exc:
        raise AttachmentTextExtractError("text_extract_failed", str(exc)) from exc
    finally:
        doc.close()


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency is installed in project env
        raise AttachmentTextExtractError("text_extract_unavailable", str(exc)) from exc

    try:
        doc = Document(str(path))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))
        return "\n\n".join(paragraphs + table_rows)
    except Exception as exc:
        raise AttachmentTextExtractError("text_extract_failed", str(exc)) from exc


def _finish(
    text: str,
    source_type: str,
    *,
    max_chars: int,
    pages_read: int = 0,
) -> ExtractedAttachmentText:
    cleaned = clean_text(text or "")
    original_chars = len(cleaned)
    truncated = original_chars > max_chars
    if truncated:
        cleaned = cleaned[:max_chars].rstrip()
    if not cleaned:
        raise AttachmentTextExtractError("empty_description")
    return ExtractedAttachmentText(
        text=cleaned,
        source_type=source_type,
        truncated=truncated,
        chars=len(cleaned),
        original_chars=original_chars,
        pages_read=pages_read,
    )


def _suffix(path: Path, name: str) -> str:
    candidate = Path(name).suffix if name else ""
    return (candidate or path.suffix).lower()


def _mime_type(path: Path, explicit: str, name: str) -> str:
    if explicit:
        return explicit.split(";", 1)[0].strip().lower()
    guessed = mimetypes.guess_type(name or str(path))[0]
    return (guessed or "").lower()
